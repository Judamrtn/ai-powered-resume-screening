"""
Resume Parser — supports both PDF and DOCX formats.
"""
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document


def parse_pdf(filepath: str) -> str:
    """Extract raw text from a PDF file."""
    return extract_pdf_text(filepath) or ""


def parse_docx(filepath: str) -> str:
    """
    Extract raw text from a DOCX file.
    Reads paragraphs AND table cells (many resumes use tables for layout).
    """
    doc   = Document(filepath)
    parts = []

    # Paragraphs (preserves top-to-bottom reading order for body text)
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables — many resume templates put contact info or skills in tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def parse_resume(filepath: str, filename: str) -> str:
    """
    Universal entry point — detects format from filename extension
    and routes to the correct parser.
    """
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return parse_pdf(filepath)
    elif lower.endswith(".docx"):
        return parse_docx(filepath)
    elif lower.endswith(".doc"):
        raise ValueError(
            "Legacy .doc format is not supported — please convert to .docx or PDF."
        )
    else:
        raise ValueError(f"Unsupported file format: {filename}")